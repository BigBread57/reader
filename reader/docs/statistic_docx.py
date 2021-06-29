from django.conf import settings
from django.core.files import File
from docx import Document
from docx.enum.section import WD_ORIENTATION
from datetime import datetime

from accountBd.models import Profile, User, FileRepository, Project


def journal(document_name):
    """
    Функция для создания шаблона документа для инструктожа

    :return: Шаблон документа для инструктажа
    """

    document = Document()

    # Для альбомногй ориентации
    section = document.sections[-1]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    document.add_heading('Список личного состава', 0)

    table = document.add_table(rows=1, cols=7)
    table.style = 'Table Grid'

    c = table.rows[0].cells
    c[0].text = 'Дата'
    c[1].text = 'Должность и звание'
    c[2].text = 'ФИО'
    c[3].text = 'Вид инструктажа'
    c[4].text = 'Подпись инструктирующего'
    c[5].text = 'Подпись инструктируемого'
    c[6].text = 'Примечание'

    document.save(f'media/docs/{document_name}')


def event_statistic(list_user, document_name):
    """
    Функция для внесения записей в журнал

    :param list_user: Список с пользователями
    :param document_name: Название файла
    :return: Добавляет список пользоваетелй в таблицу, которая создается в функции journal()
    """

    document = Document(f'media/docs/{document_name}')

    for user in list_user:
        table = document.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        c = table.rows[0].cells
        c[0].text = f'{user[0]}'
        c[1].text = f'{user[1]}\n{user[2]}'
        c[2].text = f'{user[3]}'
        c[3].text = 'целевой'
        c[4].text = ''
        c[5].text = ''
        c[6].text = ''

    document.save(f'media/docs/{document_name}')


def user_statistic(time_schedule, time_work, real_overtime,
                   time_of_respectful_absence_fact, time_of_not_respectful_absence_fact,
                   percent_time_work, percent_time_of_respectful_absence,
                   percent_time_of_not_respectful_absence,
                   position, rank, full_name, project, document_name):
    """
    Функция формирует сводный отчет по одному оператору

    :param time_schedule: Время работы по графику (плану)
    :param time_work: Реальное время работы
    :param real_overtime: Время переработок
    :param time_of_respectful_absence_fact: Уважительное время отсутствия
    :param time_of_not_respectful_absence_fact: Не уважительное время отсутствия
    :param percent_time_work: Процент рабочего времени
    :param percent_time_of_respectful_absence: Процент уважительного отсутствия
    :param percent_time_of_not_respectful_absence: Процент не уважительного отсутствия
    :param position: Должность оператора
    :param rank: Звание оператора
    :param full_name: Полное имя оператора
    :param project: Проект за которым закреплен оператор
    :param document_name: Название докмуента куда сохранять информацию
    """

    document = Document(f'media/docs/{document_name}')
    document.add_paragraph(f'Отчет по деятельности: {str(position).lower()} {str(rank).lower()} {full_name}\n\n',
                           style='Heading 1')

    table = document.add_table(rows=1, cols=2)

    table.style = 'Medium List 1 Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Название показателя'
    hdr_cells[1].text = 'Данные'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Время работы по плану '
    row_cells[1].text = str(time_schedule)

    row_cells = table.add_row().cells
    row_cells[0].text = 'Время фактической работы '
    row_cells[1].text = str(time_work)

    row_cells = table.add_row().cells
    row_cells[0].text = 'Время фактической переработки '
    row_cells[1].text = str(real_overtime)

    row_cells = table.add_row().cells
    row_cells[0].text = 'Время фактических уважительных прогулов '
    row_cells[1].text = str(time_of_respectful_absence_fact)

    row_cells = table.add_row().cells
    row_cells[0].text = 'Время фактических не уважительных прогулов '
    row_cells[1].text = str(time_of_not_respectful_absence_fact)

    row_cells = table.add_row().cells
    row_cells[0].text = 'Процент фактической работы в общем количестве времени '
    row_cells[1].text = str(percent_time_work)

    row_cells = table.add_row().cells
    row_cells[0].text = 'Процент уважительных прогулов в общем количестве времени '
    row_cells[1].text = str(percent_time_of_respectful_absence)

    row_cells = table.add_row().cells
    row_cells[0].text = 'Процент не уважительных прогулов в общем количестве времени '
    row_cells[1].text = str(percent_time_of_not_respectful_absence)

    document.add_paragraph(f'\n\n{position} {full_name} находится на проекте {project}')

    # Разрыв старницы
    document.add_page_break()

    document.save(f'media/docs/{document_name}')


# Функция для получения docx файла по одному пользователю
def user_statistic_report(serializer, user_id, document_name):
    """
    Функция позволяет привести данные к необходимому формату и пересать в функцию
     по формированию документа по одному пользователю.

    :param serializer: Получаем все данные из сериализатора
    :param user_id: id пользователя по которому необходима статистика
    :param document_name: Название документа
    """

    document = Document()
    document.save(f'media/docs/{document_name}')

    user_statistic(
        serializer['time_schedule'],
        serializer['time_work'],
        serializer['real_overtime'],
        serializer['time_of_respectful_absence_fact'],
        serializer['time_of_not_respectful_absence_fact'],
        serializer['percent_time_work'],
        serializer['percent_time_of_respectful_absence'],
        serializer['percent_time_of_not_respectful_absence'],
        Profile.objects.get(user_id=user_id).get_position_display(),
        Profile.objects.get(user_id=user_id).get_rank_display(),
        serializer['full_name'],
        serializer['project'], document_name
    )


# Функция для получения docx файла по всем пользователям
def all_users_statistic_report(queryset, document_name):
    """
    Функция отправляет обработанную информацию в функцию,
    которая формрует сводный отчет по всем операторам

    :param queryset: Получаем queryset, который содержит всех операторов
    :param document_name: Получаем название документа, куда необходимо сохранить отчет
    """

    document = Document()
    document.save(f'media/docs/{document_name}')
    profiles = list(queryset)

    for profile in profiles:
        user_statistic(
            profile.time_schedule,
            profile.time_work,
            profile.real_overtime,
            profile.time_of_respectful_absence_fact,
            profile.time_of_not_respectful_absence_fact,
            profile.time_work / profile.time_schedule * 100,
            profile.time_of_respectful_absence_fact / profile.time_schedule * 100,
            profile.time_of_not_respectful_absence_fact / profile.time_schedule * 100,
            profile.get_position_display(),
            profile.get_rank_display(),
            profile.user.get_full_name(),
            profile.project,
            document_name
        )


def project_statistic(time_schedule, time_work, real_overtime,
                      time_of_respectful_absence_fact, time_of_not_respectful_absence_fact,
                      percent_time_work, percent_time_of_respectful_absence,
                      percent_time_of_not_respectful_absence,
                      name, profiles, document_name):
    """
    Функция формирует сводный отчет по одному оператору

    :param time_schedule: Время работы по графику (плану)
    :param time_work: Реальное время работы
    :param real_overtime: Время переработок
    :param time_of_respectful_absence_fact: Уважительное время отсутствия
    :param time_of_not_respectful_absence_fact: Не уважительное время отсутствия
    :param percent_time_work: Процент рабочего времени
    :param percent_time_of_respectful_absence: Процент уважительного отсутствия
    :param percent_time_of_not_respectful_absence: Процент не уважительного отсутствия
    :param name: Название проекта
    :param profiles: Сисок профилей пользователя
    :param document_name: Название докмуента куда сохранять информацию
    """

    # Переменная для запоминания всех ФИО операторов.
    str_users = ''
    document = Document(f'media/docs/{document_name}')
    document.add_paragraph(f'Отчет по проекту: {str(name)}\n\n',
                           style='Heading 1')

    if type(profiles) is list:
        for profile in profiles:
            str_users += profile['full_name'] + ', '
    else:
        for profile in profiles:
            str_users += profile.user.get_full_name() + ', '

    document.add_paragraph(f'Операторы, которые закреплены за проектом: {str_users[:-2]}\n\n',
                           style='Heading 1')

    table = document.add_table(rows=1, cols=2)

    table.style = 'Medium List 1 Accent 1'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Название показателя'
    hdr_cells[1].text = 'Данные'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Время работы по плану '
    row_cells[1].text = str(time_schedule)

    row_cells = table.add_row().cells
    row_cells[0].text = 'Время фактической работы '
    row_cells[1].text = str(time_work)

    row_cells = table.add_row().cells
    row_cells[0].text = 'Время фактической переработки '
    row_cells[1].text = str(real_overtime)

    row_cells = table.add_row().cells
    row_cells[0].text = 'Время фактических уважительных прогулов '
    row_cells[1].text = str(time_of_respectful_absence_fact)

    row_cells = table.add_row().cells
    row_cells[0].text = 'Время фактических не уважительных прогулов '
    row_cells[1].text = str(time_of_not_respectful_absence_fact)

    row_cells = table.add_row().cells
    row_cells[0].text = 'Процент фактической работы в общем количестве времени '
    row_cells[1].text = str(percent_time_work)

    row_cells = table.add_row().cells
    row_cells[0].text = 'Процент уважительных прогулов в общем количестве времени '
    row_cells[1].text = str(percent_time_of_respectful_absence)

    row_cells = table.add_row().cells
    row_cells[0].text = 'Процент не уважительных прогулов в общем количестве времени '
    row_cells[1].text = str(percent_time_of_not_respectful_absence)

    # Разрыв старницы
    document.add_page_break()

    document.save(f'media/docs/{document_name}')


# Функция для получения docx файла по одному проекту
def project_statistic_report(serializer, document_name):
    """
    Функция позволяет привести данные к необходимому формату и пересать в функцию
    по формированию документа по одному проекту.

    :param serializer: Получаем все данные из сериализатора
    :param document_name: Название документа
    """

    document = Document()
    document.save(f'media/docs/{document_name}')

    project_statistic(
        serializer['time_schedule'],
        serializer['time_work'],
        serializer['real_overtime'],
        serializer['time_of_respectful_absence_fact'],
        serializer['time_of_not_respectful_absence_fact'],
        serializer['percent_time_work'],
        serializer['percent_time_of_respectful_absence'],
        serializer['percent_time_of_not_respectful_absence'],
        serializer['name'],
        serializer['profiles'],
        document_name
    )


# Функция для получения docx файла по всем проектамм
def projects_statistic_report(queryset, document_name):
    """
    Функция отправляет обработанную информацию в функцию,
    которая формрует сводный отчет по всем проектам

    :param queryset: Получаем queryset, который содержит все проекты
    :param document_name: Получаем название документа, куда необходимо сохранить отчет
    """

    document = Document()
    document.save(f'media/docs/{document_name}')
    projects = list(queryset)

    for project in projects:
        project_statistic(
            project.time_schedule,
            project.time_work,
            project.real_overtime,
            project.time_of_respectful_absence_fact,
            project.time_of_not_respectful_absence_fact,
            project.time_work / project.time_schedule * 100,
            project.time_of_respectful_absence_fact / project.time_schedule * 100,
            project.time_of_not_respectful_absence_fact / project.time_schedule * 100,
            project.name,
            project.profiles.all(),
            document_name
        )
